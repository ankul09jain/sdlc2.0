# import openai
import markdown2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# import google.generativeai as genai
import gradio as gr
from bs4 import BeautifulSoup
import boto3
import os
import requests
from google import genai
from google.genai import types
from botocore.exceptions import NoCredentialsError, ClientError
from dotenv import load_dotenv
import time
import traceback

# Load environment variables from a .env file
load_dotenv()

# Retrieve API keys and client secret from environment variables
openai_api_key = os.getenv("OPENAI_API_KEY")
gemini_api_key = os.getenv("GEMINI_API_KEY")
client_secret = os.getenv("CLIENT_SECRET")
client_id = os.getenv("CLIENT_ID")
tenant_id = os.getenv("TENANT_ID")
user_id = os.getenv("USER_ID")

def generate_ai_response_gemini(prompt, context=None):
    client = genai.Client(
        api_key=gemini_api_key,
    )
    files = []
    if context:
        for file in context:
            files.append(client.files.upload(file=file))
    model = "gemini-2.5-flash-preview-04-17"
    contents = [
        types.Content(
            role="user",
            parts=[
                *[
                    types.Part.from_uri(
                        file_uri=file.uri,
                        mime_type=file.mime_type,
                    ) for file in files
                ],
                types.Part.from_text(text=prompt),
            ],
        ),
    ]
    generate_content_config = types.GenerateContentConfig(
        response_mime_type="text/plain",
    )

    input_token_count = client.models.count_tokens(
        model=model,
        contents=contents,
    ).total_tokens
    print(f"\nInput tokens: {input_token_count}")



    result = ""
    output_token = 0

    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        result+=chunk.text
        output_token+= client.models.count_tokens(model=model,contents=chunk.text).total_tokens

    print(f"Output tokens: {output_token}")

    return result



def get_access_token(client_id, client_secret, tenant_id):
    url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    data = {
        "client_id": client_id,
        "client_secret": client_secret,
        "scope": "https://graph.microsoft.com/.default",
        "grant_type": "client_credentials"
    }
    response = requests.post(url, data=data)
    response.raise_for_status()
    return response.json()["access_token"]


def handle_onedrive_file(
    access_token,
    item_path,
    local_file_path=None,
    mode="download"
):
    """
    Handle OneDrive file operation (upload or download).

    Parameters:
    - access_token: Microsoft Graph access token
    - user_id: Target OneDrive user's email or object ID
    - item_path: Path in OneDrive (e.g., 'Documents/report.docx')
    - local_file_path: Local file path (used for upload or to save download)
    - mode: "download" or "upload"
    """
    print(item_path)
    headers = {"Authorization": f"Bearer {access_token}"}
    url = f"https://graph.microsoft.com/v1.0/users/{user_id}/drive/root:/{item_path}:/content"

    try:
        if mode == "download":
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            save_path = local_file_path if local_file_path else item_path.split("/")[-1]
            with open(save_path, "wb") as f:
                f.write(response.content)
            print(f"Downloaded to: {save_path}")

        elif mode == "upload":
            if not local_file_path:
                raise ValueError("local_file_path is required for upload")
            headers["Content-Type"] = "application/octet-stream"
            with open(local_file_path, "rb") as f:
                data = f.read()
            response = requests.put(url, headers=headers, data=data)
            response.raise_for_status()
            print(f"Uploaded: {item_path}")

        else:
            raise ValueError("Mode must be 'upload' or 'download'")
    except Exception as e:
        traceback.print_exc()


def generate_ai_response_openai(prompt):
    # Call OpenAI API to generate markdown content
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that writes software proposal documents in markdown."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2000,
        temperature=0.7,
    )
    return response['choices'][0]['message']['content']

def upload_to_s3(file_path, filename, bucket_name="sdlc2.0-artifacts", s3_key="discovery_phase"):

    # Initialize a session using Amazon S3
    s3 = boto3.client('s3')

    try:
        s3_key = s3_key + "/" + filename
        # Upload the file
        s3.upload_file(file_path, bucket_name, s3_key)
        print(f"File {file_path} uploaded to {bucket_name}/{s3_key}")

        # Generate a presigned URL for the uploaded file
        url = s3.generate_presigned_url('get_object',
                                        Params={'Bucket': bucket_name, 'Key': s3_key},
                                        ExpiresIn=3600)  # URL expires in 1 hour
        return url

    except FileNotFoundError:
        print(f"The file {file_path} was not found.")
        return None
    except NoCredentialsError:
        print("Credentials not available.")
        return None
    except ClientError as e:
        print(f"Client error: {e}")
        return None


def markdown_to_docx(doc_type, markdown_text, output_path, project_name):
    # Convert markdown to HTML
    html = markdown2.markdown(markdown_text)
    # Create a new Word document
    if doc_type=="s":
        template_path = "SRSTemplate.docx"
    elif doc_type=="p":
        template_path = "template.docx"
    doc = Document(template_path)
    
    # Replace "{Project Name}" in the first few paragraphs (likely cover page)
    for i, para in enumerate(doc.paragraphs[:10]):
        if "{Project Name}" in para.text:
            print(f"Found '{para.text}' on the cover page - replacing with '{project_name}'")
            para.text = para.text.replace("{Project Name}", project_name)
    
    # Add a header with the project name 
    section = doc.sections[-1]
    header = section.header
    header_paragraph = header.paragraphs[0]
    if doc_type == "s":
        header_paragraph.text = f"{project_name} Software Requirements Specifications"
    elif doc_type == "p":
        header_paragraph.text = f"{project_name}: Proposal Doc"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add content (simple approach: strip HTML tags, or use a library for better conversion)
    soup = BeautifulSoup(html, "html.parser")
    
    # Process only direct children of the body to avoid duplication
    # Get all top-level elements (direct children of body or html)
    if soup.body:
        elements = soup.body.find_all(recursive=False)
    else:
        elements = soup.find_all(recursive=False)
    
    for element in elements:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            doc.add_heading(element.get_text(), level=5)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):  
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):  
                doc.add_paragraph(li.get_text(), style='List Number')

        elif element.name == 'p':
            doc.add_paragraph(element.get_text())

    # Save the document
    doc.save(output_path)
    if doc_type=="p":
        file_name = project_name+"_Proposal.docx"
    else:
        file_name = "SRS_"+project_name+".docx"
    download_link = upload_to_s3(output_path, file_name)
    return download_link

def generate_srs_via_gradio(project_name, client_name, high_level_scope,transcript_list, artifact_list, cloud_provider):
    """
    Handles SRS generation based on inputs from the Gradio UI.
    """

    try:
        # Construct the user prompt using the global template

        user_prompt = """
You are a professional AI assistant tasked with generating a System Requirements Specification (SRS) document in markdown format. The document should accurately capture the requirements for the system based on provided inputs. Maintain a professional, clear, precise, and unambiguous tone suitable for technical documentation.

Use the following parameters to customize the SRS:

* Client Name: {client_name}
* Project/System Name: {project_name}
* High-Level System Scope: {high_level_scope}
* Input Sources (e.g., Transcripts, Artifacts): {input_sources}
* Cloud Provider: {cloud_provider}

Your task is to analyze the information contained within the Input sources to extract, synthesize, and document the system requirements for the {project_name} being developed for {client_name}. Structure the document based on the sections and subsections outlined below, drawing inspiration from the provided example SRS document structure. You are also authorized to create new sections or subsections if the analyzed input sources reveal requirements or context that do not fit neatly into the predefined structure but are essential for a complete SRS, guided by the overall project scope and goals.

Here is the recommended structure and guidance for each section of the SRS, inspired by the example document, with added requirements for flows, UI descriptions, validations, and architecture:

**1. Introduction:**

* **Guidance:** Provide a foundational overview of the document and the project.
* **Content to include:**
    * 1.1 Purpose: Explain the purpose of this SRS document.
    * 1.2 Project Scope: Detail the boundaries and key features included in the {project_name} based on the High-Level System Scope and input sources.
    * 1.3 Intended Audience: Specify who this document is for (e.g., developers, testers, stakeholders).

**2. Overall Description:**

* **Guidance:** Describe the general context and factors influencing the system.
* **Content to include:**
    * 2.1 Product Perspective: Explain how the {project_name} fits into the larger picture or interacts with other systems.
    * 2.2 Operating Environment: Describe the environments where the system will operate (e.g., mobile platforms, web browsers, servers).
    * 2.3 Assumption & Dependency: List factors assumed to be true and external dependencies the project relies on (e.g., third-party services, data availability).

**3. System Architecture:**

* **Guidance:** Provide a high-level description of the system's architecture, focusing on its main components and how they interact. Specify the **{cloud_provider}** as the chosen cloud platform for deployment and relevant services.
* **Content to include:** Architectural style (e.g., serverless, microservices), key components (e.g., mobile app, backend services, database, third-party integrations), and how they communicate. Mention the role of the specified **{cloud_provider}** and relevant services (e.g., compute, database, storage, AI/ML services).

**4. User Flows and Interface Details:**

* **Guidance:** Describe the step-wise user flows through the system for key processes. **For each user journey or major process, structure the details clearly by screen name using a level 4 markdown heading (####).** Under each screen name heading, provide a concise description (under 200 words), a description of validations/error handling (under 100 words), and a placeholder for the corresponding wireframe/UI image. Use bullet points for the Description and Validations/Error Handling under the screen name heading. Extract these flows and details from the Input sources.
* **Content to include:**
    * For each major user journey or process (e.g., User Onboarding Flow, Login Process Flow, \[Feature Name\] Workflow):
        * Provide a brief introduction to the flow.
        * **For each screen within the flow add a new subsection for the Name of the screen and then provide below details
            * **Description:** \[Concise textual description (under 150 words) of the screen's layout, main components (buttons, input fields, labels, images, etc.), and its purpose in the flow.\]
            * **Validations/Error Handling:** \[Concise description (under 100 words) of any specific input validations, error messages, or handling logic related to this screen.\]
            * **Placeholder:** `[Placeholder: Wireframe/UI for Screen Name]`

**5. Specific Requirements:**

* **Guidance:** Detail the functional and non-functional requirements. Organize these logically. Each requirement should be clear, unambiguous, and ideally testable. Extract these details by analyzing the Input sources.
* **Content to include:**
    * **Functional Requirements:** Describe what the system *must do*. Organize these logically (e.g., by feature, user role, or system mode). Use clear, action-oriented language. Each requirement should ideally be a distinct point.
    * **Non-Functional Requirements:** Describe *how* the system should perform or qualities it must possess. Include categories like:
        * Performance Requirements (e.g., response time, throughput)
        * Security Requirements (e.g., access control, data protection)
        * Usability Requirements (e.g., ease of use, user interface standards)
        * Reliability Requirements (e.g., availability, error handling)
        * Maintainability Requirements (e.g., ease of modification, adherence to standards)
        * Scalability Requirements (e.g., ability to handle increased load)
        * Environmental Requirements (e.g., operating system, hardware)

**(Continue adding sections for other major components/feature areas as identified from the example document structure and the input sources, e.g., Plan Subscription Details, Bank Account Integration Details, Credit Monitoring Service Integration Details, etc. Integrate specific functional requirements related to these features within these sections or cross-reference them from the "Specific Requirements" section.)**

\[Other Relevant Sections from Example, e.g., Terms & Conditions, Privacy Policy, References, Out of Scope for MVP\]:

* **Guidance:** Include sections for important related information as seen in the example SRS. For sections marked "Pending" in the example, note that the details are TBD based on input sources. For "Out of Scope," clearly list items explicitly excluded from the current phase.

\[New Sections (if needed)\]:

* **Guidance:** If the Input sources contain significant requirements, constraints, or context that do not fit into the above structure, create new, appropriately titled sections or subsections to capture them. Ensure these new sections are relevant to the {project_name} and its scope.

**Conclusion:**

* **Guidance:** Provide a brief concluding statement summarizing the SRS and its role in the project.

**General Instructions for AI:**

* Generate the entire response in markdown format.
* Use clear headings and subheadings corresponding to the SRS structure outlined above. Use appropriate markdown heading levels (`#`, `##`, `###`,`####`, `#####` etc.). Ensure a consistent and logical heading hierarchy throughout the document.
* **Generate a markdown subheading for every subsection created under each main section of the document.** Use appropriate markdown heading levels for these subheadings.
* Use bullet points or numbered lists for requirements within sections, particularly in "Specific Requirements" and within the step-wise flows. **Only use bullet points or numbered lists for the lowest level of detail within a section or subsection.**
* Maintain a formal, technical, and objective tone throughout the document.
* Carefully analyze the Input sources to extract accurate, complete, and unambiguous requirements. Synthesize information from multiple sources if necessary.
* Incorporate the values provided for each parameter throughout the relevant sections of the document.
* Creatively include new sections or subsections if the input data necessitates them for a comprehensive SRS, ensuring they align with the project scope and goals.
* For UI/wireframe descriptions, provide enough detail for clarity and explicitly include the `[Placeholder: Wireframe/UI for ...]` text where the image should be inserted in the final document.
* Clearly describe validation rules and error handling within the relevant sections or flow steps.
* Do not create any specific alphanumeric identifiers for functional requirements, non-functional requirements, constraints, assumptions, or dependencies.

"""
        transcript_file_list=[]
        artifact_file_list=[]

        token = get_access_token(client_id, client_secret, tenant_id)
        # Download file for onedrive
        for file_name in transcript_list.split(","):
            print(file_name)
            unique_number = str(int(time.time() * 1000))[-6:]
            file_path="inputs/" + unique_number + file_name
            handle_onedrive_file(
                token,
                item_path=f"Documents/{file_name}",
                local_file_path=file_path,
                mode="download"
            )
            transcript_file_list.append(file_path)
        print(transcript_file_list)
        
        for file_name in artifact_list.split(","):
            unique_number = str(int(time.time() * 1000))[-6:]
            file_path="inputs/" + unique_number + file_name 
            handle_onedrive_file(
                token,
                item_path=f"Documents/{file_name}",
                local_file_path=file_path,
                mode="download"
            )
            artifact_file_list.append(file_path)
        print(artifact_file_list)

        transcript_prompt = "Summarize the client meeting transcript files, ensuring all client instructions are captured accurately."
        transcripts = generate_ai_response_gemini(transcript_prompt, context=["transcript.txt"])

        print("Transcripts summarized by AI")

        artifact_prompt = "Summarize these project discovery files, keep all the cruical application featueres and scope as is."
        artifacts = generate_ai_response_gemini(artifact_prompt, context=["transcript.txt"])

        print("Artifacts summarized by AI")

        # #Read artifacts from a hardcoded file
        # artifacts_file = "artifacts.txt"
        # try:
        #     with open(artifacts_file, "r", encoding="utf-8") as f:
        #         artifacts = f.read()
        # except FileNotFoundError:
        #     print(f"Artifacts file not found at {artifacts_file}.")
        #     artifacts = ""

        # # Read transcripts from a hardcoded file with encoding fallback
        # transcripts_file = "transcript.txt"
        # try:
        #     with open(transcripts_file, "r", encoding="utf-8") as f:
        #         transcripts = f.read()
        # except UnicodeDecodeError:
        #     with open(transcripts_file, "r", encoding="latin-1") as f:
        #         transcripts = f.read()
        # except FileNotFoundError:
        #     print(f"Transcripts file not found at {transcripts_file}.")
        #     transcripts = ""

        input_sources = f"Client Meeting Transcripts:\n{transcripts}\n\nAdditional Artifacts:\n{artifacts}"


        user_prompt = user_prompt.replace(
            "{project_name}", project_name
        ).replace(
            "{client_name}", client_name
        ).replace(
            "{high_level_scope}", high_level_scope
        ).replace(
            "{input_sources}", input_sources
        ).replace(
            "{cloud_provider}", cloud_provider
        )

        print(f"\n[Gradio] Generating SRS document for: {project_name}")
        print("\n==============Complete User prompt============\n" + user_prompt)
        markdown_content = generate_ai_response_gemini(user_prompt)

        # # Read SRS template from file
        # markdown_file = os.path.join(BASE_DIR, "SRS_PrimetimePharmacy.md")
        # try:
        #     with open(markdown_file, "r", encoding="utf-8") as f:
        #         markdown_content = f.read()
        # except FileNotFoundError:
        #     print(f"Template file not found at {markdown_file}")
        
        # Save markdown content to .md file
        md_file_gradio = f"outputs/SRS_{project_name}.md"
        with open(md_file_gradio, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        print(f"\n[Gradio] Markdown file saved to {md_file_gradio}")

        # Convert markdown to .docx file
        output_file_gradio = f"outputs/SRS_{project_name}.docx"
        download_link = markdown_to_docx("s", markdown_content, output_file_gradio, project_name)
        print(f"\n[Gradio] Word document saved to {output_file_gradio}")

        return f"SRS document generated successfully for '{project_name}'.\n Download the SRS document for this link: {download_link}"
    except Exception as e:
        print(f"[Gradio] Error during SRS generation: {e}")
        traceback.print_exc()
        return f"Error generating SRS: {e}"


def generate_proposal_via_gradio(industry_name,client_name, client_request, tech_stack, average_hourly_rate, estimation_notes):
    try:
        print(f"\n[Gradio Proposal] Received request for: {client_name}")
        

        prompt = f"""
        You are a professional AI assistant tasked with generating a project proposal document in markdown format. The proposal should follow a standard structure and incorporate the specific project details provided through parameters. For sections like responsibilities, timeline, and cost, generate content creatively based on the project scope described by the other parameters. Maintain a professional, clear, and persuasive tone, allowing for creativity while ensuring accuracy and adherence to the structure.

        Use the following parameters to customize the proposal:
        - Industry Name: {industry_name}
        - Client Name: {client_name}
        - Client Request: {client_request}
        - Preferred Technology Stack: {tech_stack}
        - Additional Notes/Disclaimers for Effort Estimation: {estimation_notes}

        Here is the structure and guidance for each section:

        **Document Title:**
        * **Guidance:** Create a concise title for the proposal, incorporating the Industry Name and Client Name.

        **Executive Summary:**
        * **Guidance:** Provide a brief overview. Summarize the problem addressed by the Client Request, the proposed solution, the scope of work (creatively define typical responsibilities based on the project type and client/your roles), and key benefits for the Client Name in the {industry_name} sector. Keep the content within 100 words.

        **Project Overview:**
        * **Guidance:** Describe the project context based on the Client Request, the objectives, and the Client Name's goals. Explain the value proposition of the proposed solution for their specific needs in the {{industry_name}} sector.

        **Project Approach:**
        * **Guidance:** Outline a typical methodology and phases for a project addressing the Client Request using the Preferred Technology Stack. Describe the key activities in each phase and creatively estimate a plausible timeline based on the project scope. Emphasize collaboration and integration. Keep the content within 150 words. Structure the content using bullet points in a way that is easy to understand and follow.

        **Proposed Solution Details:**
        * **Guidance:** Provide detailed descriptions of the core components, modules, or features of the proposed solution designed to address the Client Request. Structure the solution in bullet points in a way that is easy to understand and follow.
        * **Content to include:** Describe the main modules and features of the solution and keep it concise.

        **Proposed Technology Stack & Tools:**
        * **Guidance:** Think about the {client_request} core modules and features and briefly describe the technologies, frameworks, and tools you propose to use, incorporating the Preferred Technology Stack. Mention the rationale for selection based on the Client Request and {industry_name} needs. Keep the content within 150 words. Structure the content in bullet points in a way that is easy to understand and follow.

        **Assumptions & Dependencies:**
        * **Guidance:** List any assumptions made about the Client Name's environment, data, or resources related to the Client Request. List any dependencies typically required for project success, including dependencies related to client responsibilities (creatively defined).

        **Effort Estimation:**
        * **Guidance:** {estimation_notes} \n Present the timeline and cost. Include any Additional Notes/Disclaimers for Effort Estimation. Keep the final presented content concise (under 100 words). Structure the content in bullet points in a way that is easy to understand and follow. 

        **Conclusion:**
        * **Guidance:** Write a concluding statement reinforcing commitment to the Client Name and expressing enthusiasm for collaboration on the project addressing the Client Request.

        **General Instructions for AI:**
        * Generate the entire response in markdown format.
        * Use clear headings and subheadings, bullet points and formatted text to make the content more readable.
        * Use bullet points instead of numbers where appropriate (e.g., for solution details, technology stack, approach assumptions, dependencies).
        * Maintain a professional tone.
        * Incorporate the values provided for each parameter throughout the relevant sections of the proposal.
        * Creatively generate content for team responsibilities, client responsibilities, and estimated timeline based on the context provided by the other parameters and the general nature of the project described in the Client Request and Solution Details.
        * **For Effort Estimation, estimate the number of resources needed for Project Management, Quality Assurance, Developers, and DevOps based on the project scope and timeline.**
        * **Calculate the estimated cost range by multiplying the estimated total hours (based on timeline and resource allocation) by the provided Average Hourly Rate {average_hourly_rate}.**
        * Ensure to keep the proposal concise and to the point and avoid any duplication or fluff.
        * Don't mention the word count in the proposal.
        * Ensure the estimated cost range reflects a plausible calculation based on the estimated resources, timeline, and average hourly rate.

        """

        # Assuming generate_proposal_gemini is available in the scope (e.g., imported)
        markdown_content = generate_ai_response_gemini(prompt)
        print(f"[Gradio Proposal] Markdown content generated for proposal.")

        # Sanitize project_name for filename to avoid issues with special characters
        safe_project_name = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in client_name).replace(' ', '_')
        
        md_file_proposal = f"outputs/Proposal_{safe_project_name}.md"
        with open(md_file_proposal, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        print(f"\n[Gradio Proposal] Markdown file saved to {md_file_proposal}")

        output_file_proposal = f"outputs/Proposal_{safe_project_name}.docx"
        # Assuming markdown_to_docx is available in the scope (e.g., imported)
        # and uses a template.docx from "Desktop/template.docx" as per genai.py context
        download_link = markdown_to_docx("p", markdown_content, output_file_proposal, client_name)
        print(f"\n[Gradio Proposal] Word document saved to {output_file_proposal}")

        return f"Proposal document generated successfully for '{client_name}'.\nProposal doc Download Link: {download_link}"
    except Exception as e:
        print(f"[Gradio Proposal] Error during proposal generation: {e}")
        # For more detailed debugging, one might add:
        # import traceback
        # traceback.print_exc()
        return f"Error generating proposal: {e}"


if __name__ == "__main__":

    print("\n=== SRS & Proposal Document Generator ===\n")

    # markdown_content =""
    # #Save markdown content to .md file
    # md_file = "/Users/ankuljain/Desktop/SRS_Primetime Pharmacy.md"
    # with open(md_file, "r") as f:
    #     markdown_content = f.read()
    # output_file = "/Users/ankuljain/Desktop/SRS_Primetime Pharmacy.docx"
    # markdown_to_docx(markdown_content, output_file, project_name="Primetime Pharmacy")
    # print(f"\nSRS document saved to {output_file}")

    # Define the Gradio interface
    # This interface will be launched if the surrounding code block is executed.
    srs_interface = gr.Interface(
        fn=generate_srs_via_gradio,
        inputs=[
            gr.Textbox(label="Project Name", placeholder="Enter the project name"),
            gr.Textbox(label="Client Name", placeholder="Enter the client name"),
            gr.Textbox(label="High-Level System Scope", placeholder="Enter the project goal)", lines=3),
            gr.Textbox(label="Project Transcript Filename", placeholder="transcript1.pdf, transcript2.txt", info="Provide the filename present in the above onedrive location; \nSupported File types: .pdf, .txt"),
            gr.Textbox(label="Project Artifact Filename", placeholder="artifac1.txt, artifac2.pdf", info="Just Provide the filename present in the above onedrive location;  \nSupported File types: .pdf, .txt"),
            gr.Dropdown(
                label="Cloud PLatform",
                choices=["AWS", "Azure", "GCP"],
                value="AWS",  # Default value
                interactive=True
            ),
        ],
        outputs=gr.Textbox(label="Generation Status and File Paths", lines=10),
        title="Software Requirements Specification (SRS) Generator",
        description="Enter project details to generate an SRS document. Please note this app at present only has access to download transcripts and artifacts file from this [onedrive location](https://matellioco-my.sharepoint.com/:f:/g/personal/ankul_jain_matellio_net/Esywc2R7S2xHlEm3ScB8_N8BvKJAy6kKc_fkgw62YGASQw?e=pGRLMT)"
    )

    # Define the handler function for proposal generation
    
    # Define the Gradio interface for proposal generation
    estimation_notes = """Creatively estimate a plausible Timeline range (e.g., in weeks) for a project of this scope, considering the Client Request and Preferred Technology Stack. **Estimate the required resources across key roles:** Project Management, Quality Assurance, Developers (consider different levels if appropriate based on scope), and DevOps. **Calculate the estimated cost range based on the estimated timeline, resource allocation, and the provided Average Hourly Rate.** Consider Agile and fastrack delivery with No Delivery phase """
    proposal_interface = gr.Interface(
        fn=generate_proposal_via_gradio,
        inputs=[
            gr.Textbox(label="Industry Name"),
            gr.Textbox(label="Client Name"),
            gr.Textbox(label="Client Request", placeholder="What is the project goal?"),
            gr.Textbox(label="Tech Stack", placeholder = "eg. AWS, Python, OpenAI GPT4o, ReactJS"),
            gr.Textbox(label="Average Hourly Rate($)"),            
            gr.Textbox(label="Estimation Notes(Please review)", value=estimation_notes, lines=5, autofocus=True),

        ],
        outputs=gr.Textbox(label="Generation Status and File Paths", lines=10),
        title="Client Proposal Generator",
        description="Enter project details to generate a proposal document."
    )

    # Combine the existing SRS interface and the new Proposal interface into a TabbedInterface
    # The colleague will need to change the .launch() call from srs_interface.launch() to 
    tabbed_interface = gr.TabbedInterface(
        [srs_interface, proposal_interface], 
        tab_names=["SRS Generator", "Proposal Generator"]
    )


    print("\nGradio interface launching...")
    tabbed_interface.launch(server_name="0.0.0.0", server_port=7871)





