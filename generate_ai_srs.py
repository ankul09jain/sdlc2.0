import markdown2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from bs4 import BeautifulSoup, Tag, NavigableString
import boto3
import os
import requests
from botocore.exceptions import NoCredentialsError, ClientError
from dotenv import load_dotenv
import time
import traceback
import base64 # Added for image encoding
import json # Added for JSON handling in API calls
from docx.shared import Inches
from openai import OpenAI
import time
import zipfile
import tempfile
# import gradio as gr
import streamlit as st
from pathlib import Path
import re
from PyPDF2 import PdfReader
import io
import pandas as pd


# Load environment variables from a .env file
load_dotenv()

# Retrieve API keys and client secret from environment variables
openai_api_key_srs = os.getenv("OPENAI_API_KEY_SRS_GEN")
openai_api_key_proposal = os.getenv("OPENAI_API_KEY_PROPOSAL_GEN")
gemini_api_key = os.getenv("GEMINI_API_KEY")
client_secret = os.getenv("CLIENT_SECRET")
client_id = os.getenv("CLIENT_ID")
tenant_id = os.getenv("TENANT_ID")
user_id = os.getenv("USER_ID")

# DigitalOcean Spaces credentials (ensure these are set in your environment or .env)
do_spaces_region = os.getenv("DO_SPACES_REGION")
do_spaces_endpoint = os.getenv("DO_SPACES_ENDPOINT")
do_spaces_access_key = os.getenv("DO_SPACES_ACCESS_KEY")
do_spaces_secret_key = os.getenv("DO_SPACES_SECRET_KEY")


def generate_ai_response_gemini(prompt, context=None):
    # Configure the API key
    genai.configure(api_key=gemini_api_key)

    uploaded_files = []
    if context:
        for file_path in context:
            # Use genai.upload_file directly
            uploaded_file = genai.upload_file(path=file_path)
            uploaded_files.append(uploaded_file)
            print(f"Uploaded file URI: {uploaded_file.uri}")

    model_name = "gemini-2.5-flash-preview-04-17" # Using the preview model as per your original code
    model = genai.GenerativeModel(model_name)

    # Build the content parts using dictionary format for robustness
    parts = []
    for file in uploaded_files:
        parts.append({
            "file_data": {
                "file_uri": file.uri,
                "mime_type": file.mime_type
            }
        })
    parts.append({"text": prompt}) # For text parts

    # Construct contents as a list of dictionaries
    contents = [
        {
            "role": "user",
            "parts": parts,
        },
    ]

    # Define generation configuration separately
    generation_config = {
        "response_mime_type": "text/plain",
    }

    # Counting tokens using model.count_tokens directly
    input_token_count = model.count_tokens(
        contents=contents,
    ).total_tokens
    print(f"\nInput tokens: {input_token_count}")

    result = ""
    output_token = 0

    # Generating content stream using model.generate_content
    for chunk in model.generate_content(
        contents=contents,
        generation_config=generation_config,
        stream=True
    ):
        result += chunk.text
        # Updated token counting to use model.count_tokens
        output_token += model.count_tokens(contents=[{"text": chunk.text}]).total_tokens

    print(f"Output tokens: {output_token}")

    # Optionally, delete uploaded files if they are temporary and no longer needed
    # for file in uploaded_files:
    #     genai.delete_file(file.name) # Use file.name for deletion
    #     print(f"Deleted file: {file.name}")

    return result


def generate_ai_response_openai(prompt,context=None):
    client = OpenAI(api_key=openai_api_key_srs)
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that writes software requirements documents in markdown."},
            {"role": "user", "content": prompt+ "\nContext:\n" + str(context)}
        ],
        max_tokens=2000,
        temperature=0.7,
    )
    return response.choices[0].message.content


def upload_to_s3(file_path, filename, bucket_name="sdlc2.0-artifacts", s3_key="discovery_phase"):
    # Initialize a session using DigitalOcean Spaces
    # Use environment variables for credentials and endpoint
    try:
        s3 = boto3.client(
            's3',
            region_name=do_spaces_region,
            endpoint_url=do_spaces_endpoint,
            aws_access_key_id=do_spaces_access_key,
            aws_secret_access_key=do_spaces_secret_key
        )
    except Exception as e:
        print(f"Error initializing S3 client for DigitalOcean Spaces: {e}")
        traceback.print_exc()
        return None

    try:
        s3_key = s3_key + "/" + filename
        # Upload the file
        s3.upload_file(file_path, bucket_name, s3_key)
        print(f"File {file_path} uploaded to {bucket_name}/{s3_key}")

        # Generate a presigned URL for the uploaded file
        url = s3.generate_presigned_url('get_object',
                                        Params={'Bucket': bucket_name, 'Key': s3_key},
                                        ExpiresIn=3600)  # URL expires in 1 hour
        return url

    except FileNotFoundError:
        print(f"The file {file_path} was not found.")
        return None
    except NoCredentialsError:
        print("DigitalOcean Spaces credentials not available. Make sure AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, DO_SPACES_REGION, and DO_SPACES_ENDPOINT are set.")
        return None
    except ClientError as e:
        print(f"Client error during upload to DigitalOcean Spaces: {e}")
        traceback.print_exc()
        return None


def markdown_to_docx(doc_type, markdown_text, output_path, project_name):
    # Convert markdown to HTML
    html = markdown2.markdown(markdown_text, extras=["tables"])
    # print(html)
    # Create a new Word document
    if doc_type=="s":
        template_path = "SRSTemplate.docx"
    elif doc_type=="p":
        template_path = "template.docx"
    doc = Document(template_path)
    
    # Replace "{Project Name}" in the first few paragraphs (likely cover page)
    for i, para in enumerate(doc.paragraphs[:50]):
        if "{Project Name}" in para.text:
            print(f"Found '{para.text}' on the cover page - replacing with '{project_name}'")
            para.clear()
            # Add the new project name with the desired hardcoded styling
            new_run = para.add_run(project_name)
            new_run.bold = True # Make it bold
            new_run.font.name = "Arial" # Set a common, clean font
            new_run.font.size = Inches(0.5) # A good, prominent size (36pt)
            new_run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00) # Decent Orange color (RGB: 255, 165, 0)
            break
    
    # Add a header with the project name 
    section = doc.sections[-1]
    header = section.header
    header_paragraph = header.paragraphs[0]
    if doc_type == "s":
        header_paragraph.text = f"{project_name} Software Requirements Specifications"
    elif doc_type == "p":
        header_paragraph.text = f"{project_name}: Proposal Doc"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add content (simple approach: strip HTML tags, or use a library for better conversion)
    soup = BeautifulSoup(html, "html.parser")
    
    # Process only direct children of the body to avoid duplication
    # Get all top-level elements (direct children of body or html)
    if soup.body:
        elements = soup.body.find_all(recursive=False)
    else:
        elements = soup.find_all(recursive=False)
    
    for element in elements:
        try:
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(),level=level)
            elif element.name == 'table':
                # Find all rows
                rows = element.find_all('tr')
                if not rows:
                    continue
                # Count columns from the first row
                cols = rows[0].find_all(['td', 'th'])
                num_cols = len(cols)
                num_rows = len(rows)
                # Add table to docx
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        # Set cell text
                        table.cell(i, j).text = cell.get_text()
                        # Optionally, bold header row
                        if i == 0 and cell.name == 'th':
                            for paragraph in table.cell(i, j).paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):  
                    doc.add_paragraph(li.get_text(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):  
                    doc.add_paragraph(li.get_text(), style='List Number')
            elif element.name == 'p':
                # print("inside Para element tag: " + str(element.name) + str(element.contents))
                p = doc.add_paragraph()
                for content in element.contents:
                    if isinstance(content, Tag):
                        # If the child is a tag, apply formatting based on its name
                        if content.name in ['strong', 'b']:
                            p.add_run(content.get_text()).bold = True
                        elif content.name in ['em', 'i']:
                            p.add_run(content.get_text()).italic = True
                        elif content.name == 'br': # Handle line breaks
                            p.add_run('\n')
                        elif content.name == 'img': 
                            img_src = content.get('src')
                            if img_src and os.path.exists(img_src):
                                try:
                                    # Set a maximum width for the image to prevent overfitting
                                    doc.add_picture(img_src, width=Inches(6))
                                    print(f"Inserted image: {img_src}")
                                except Exception as img_e:
                                    print(f"Error inserting image {img_src}: {img_e}")
                        else:
                            # Fallback for other HTML tags (e.g., <span>, <div> within a <p>)
                            # Just add their text without special formatting for now
                            p.add_run(content.get_text())
                    elif isinstance(content, NavigableString):
                        # If the child is plain text, add it as a normal run
                        p.add_run(str(content))                 
        except KeyError:
            doc.add_paragraph(element.get_text(), style='Normal')    
    # Save the document
    doc.save(output_path)
    print(f"Document saved to {output_path}")
    if doc_type=="p":
        file_name = project_name+"_Proposal.docx"
    else:
        file_name = "SRS_"+project_name+".docx"
    download_link = upload_to_s3(output_path, file_name)
    return download_link


def generate_srs_text(project_name, client_name, high_level_scope,transcript_list, artifact_list, cloud_provider):
    """
    Handles SRS generation based on inputs from the Gradio UI.
    """

    try:
        # Define prompts at the very beginning of the function
        transcript_prompt = "Summarize the client meeting transcript files, ensuring all client instructions are captured accurately."
        artifact_prompt = "Summarize these project discovery files, keep all the cruical application features and scope as is."
        transcripts = artifacts = ""
        if transcript_list:
            transcripts = generate_ai_response_openai(transcript_prompt, context=transcript_list)
            print("Transcripts summarized by AI")
        if artifact_list:
            artifacts = generate_ai_response_openai(artifact_prompt, context=artifact_list)
            print("Artifacts summarized by AI")

        input_sources = f"""Client Meeting Transcripts:\n{transcripts}\n\nAdditional Artifacts:\n{artifacts}"""


        user_prompt = f"""
You are a professional AI assistant tasked with generating a System Requirements Specification (SRS) document in markdown format. The document should accurately capture the requirements for the system based on provided inputs. Maintain a professional, clear, precise, and unambiguous tone suitable for technical documentation.

Use the following parameters to customize the SRS:

* Client Name: {client_name}
* Project/System Name: {project_name}
* High-Level System Scope: {high_level_scope}
* Input Sources (e.g., Transcripts, Artifacts): {input_sources}
* Cloud Provider: {cloud_provider}

Your task is to analyze the information contained within the Input sources to extract, synthesize, and document the system requirements for the {project_name} being developed for {client_name}. Structure the document based on the sections and subsections outlined below, drawing inspiration from the provided example SRS document structure. You are also authorized to create new sections or subsections if the analyzed input sources reveal requirements or context that do not fit neatly into the predefined structure but are essential for a complete SRS, guided by the overall project scope and goals.

Here is the recommended structure and guidance for each section of the SRS, inspired by the example document, with added requirements for flows, UI descriptions, validations, and architecture:

**1. Introduction:**

* **Guidance:** Provide a foundational overview of the document and the project.
* **Content to include:**
    * 1.1 Purpose: Explain the purpose of this SRS document.
    * 1.2 Project Scope: Detail the boundaries and key features included in the {project_name} based on the High-Level System Scope and input sources.
    * 1.3 Intended Audience: Specify who this document is for (e.g., developers, testers, stakeholders).

**2. Overall Description:**

* **Guidance:** Describe the general context and factors influencing the system.
* **Content to include:**
    * 2.1 Product Perspective: Explain how the {project_name} fits into the larger picture or interacts with other systems.
    * 2.2 Operating Environment: Describe the environments where the system will operate (e.g., mobile platforms, web browsers, servers).
    * 2.3 Assumption & Dependency: List factors assumed to be true and external dependencies the project relies on (e.g., third-party services, data availability).

**3. System Architecture:**

* **Guidance:** Provide a high-level description of the system's architecture, focusing on its main components and how they interact. Specify the **{cloud_provider}** as the chosen cloud platform for deployment and relevant services.
* **Content to include:** Architectural style (e.g., serverless, microservices), key components (e.g., mobile app, backend services, database, third-party integrations), and how they communicate. Mention the role of the specified **{cloud_provider}** and relevant services (e.g., compute, database, storage, AI/ML services).


**4. Specific Requirements:**

* **Guidance:** Detail the functional and non-functional requirements. Organize these logically. Each requirement should be clear, unambiguous, and ideally testable. Extract these details by analyzing the Input sources.
* **Content to include:**
    * **Functional Requirements:** Describe what the system *must do*. Organize these logically (e.g., by feature, user role, or system mode). Use clear, action-oriented language. Each requirement should ideally be a distinct point.
    * **Non-Functional Requirements:** Describe *how* the system should perform or qualities it must possess. Include categories like:
        * Performance Requirements (e.g., response time, throughput)
        * Security Requirements (e.g., access control, data protection)
        * Usability Requirements (e.g., ease of use, user interface standards)
        * Reliability Requirements (e.g., availability, error handling)
        * Maintainability Requirements (e.g., ease of modification, adherence to standards)
        * Scalability Requirements (e.g., ability to handle increased load)
        * Environmental Requirements (e.g., operating system, hardware)

**(Continue adding sections for other major components/feature areas as identified from the example document structure and the input sources, e.g., Plan Subscription Details, Bank Account Integration Details, Credit Monitoring Service Integration Details, etc. Integrate specific functional requirements related to these features within these sections or cross-reference them from the "Specific Requirements" section.)**

[Other Relevant Sections from Example, e.g., Terms & Conditions, Privacy Policy, References, Out of Scope for MVP]:

* **Guidance:** Include sections for important related information as seen in the example SRS. For sections marked "Pending" in the example, note that the details are TBD based on input sources. For "Out of Scope," clearly list items explicitly excluded from the current phase.

[New Sections (if needed)]:
If you are creating new sections, ensure they are relevant and appropriately named in the final markdown SRS output. Eg. for Agentic AI based Project, create a section proposing various Agents to be created.
* **Guidance:** If the Input sources contain significant requirements, constraints, or context that do not fit into the above structure, create new, appropriately titled sections or subsections to capture them. Ensure these new sections are relevant to the {project_name} and its scope.


**General Instructions for AI:**

* Generate the entire response in markdown format.
* Use clear headings and subheadings corresponding to the SRS structure outlined above. Use appropriate markdown heading levels (`#`, `##`, `###`,`####`, `#####` etc.). Ensure a consistent and logical heading hierarchy throughout the document.
* **Generate a markdown subheading for every subsection created under each main section of the document.** Use appropriate markdown heading levels for these subheadings.
* Use bullet points or numbered lists for requirements within sections, particularly in "Specific Requirements" and within the step-wise flows. **Only use bullet points or numbered lists for the lowest level of detail within a section or subsection.**
* Maintain a formal, technical, and objective tone throughout the document.
* Carefully analyze the Input sources to extract accurate, complete, and unambiguous requirements. Synthesize information from multiple sources if necessary.
* Incorporatethe values provided for each parameter throughout the relevant sections of the document.
* Clearly describe validation rules and error handling within the relevant sections or flow steps.
* Do not create any specific alphanumeric identifiers for functional requirements, non-functional requirements, constraints, assumptions, or dependencies.

"""
        print(f"\nGenerating SRS document for: {project_name}")
        # print("\n==============Complete User prompt============\n" + user_prompt)
        markdown_content = generate_ai_response_openai(user_prompt)

        return markdown_content

    except Exception as e:
        print(f"[Gradio] Error during SRS generation: {e}")
        traceback.print_exc()
        return f"Error generating SRS: {e}"


def image_to_base64(image_path):
    """
    Encodes an image file to a base64 string.
    """
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    except FileNotFoundError:
        print(f"Error: Image file not found at {image_path}")
        return None
    except Exception as e:
        print(f"Error encoding image: {e}")
        return None

def analyze_ui_mockup_with_gemini(image_path, HIGH_LEVEL_SCOPE):
    """
    Analyzes a UI mockup image using the Gemini API and generates a markdown table.
    """
    # Replace with your actual Gemini API key if not running in a Canvas environment
    # In a Canvas environment, the apiKey will be provided automatically in the fetch call.
    # api_key = os.environ.get("GEMINI_API_KEY", "")
    api_key = os.getenv("GEMINI_API_KEY")
    api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"

    base64_image = image_to_base64(image_path)
    if not base64_image:
        return "Failed to encode image."

    prompt_text = f"""
Analyze the provided UI screenshot thoroughly, considering the high-level project scope: {HIGH_LEVEL_SCOPE}. Your output should be structured into three distinct subsections:
1. **Screen Name**: First, identify the primary purpose or title of this UI screen and present it as a clear heading. Make this markdown bold line

2.  **Screen Description (Approx. 75-100 words):**
    Provide a comprehensive description of the screen. This description should cover:
    * The overall purpose of the screen and its main functionality.
    * The primary layout and organization of elements.
    * Any key workflows or user interactions facilitated by this screen.
    * Its general appearance and any notable visual cues.
    Ensure this description is around 75-100 words and is plain text with no heading.

3.  **UI Elements Markdown Table:**
    Create a markdown table detailing all interactive elements, input fields, and buttons on the screen. For each identified element, include the following columns:
    * **Field Name**: The explicit name, label, or clear identifier associated with the UI element.
    * **Type**: The specific type of UI element (e.g., 'Text Input', 'Button', 'Dropdown', 'Text Area', 'Display Area', 'Link').
    * **Description**: A concise explanation of its specific purpose and functionality within this UI context.
    * **Validation Rules**: Any visible, implied, or standard validation rules or constraints applicable to input fields (e.g., 'Required', 'Numeric Only', 'File Types: .pdf, .txt', 'Specific Format').

    Ensure the table is complete, accurately reflects all interactive elements, and strictly uses the specified column names. 
    Don't add any style or formatting to the table or headings.

The output should not include any unneccesary white or blank blocks.
    """

    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {"text": prompt_text},
                    {
                        "inlineData": {
                            "mimeType": "image/jpeg",  # Adjust if your image is PNG
                            "data": base64_image
                        }
                    }
                ]
            }
        ],
        "generationConfig": {
            "responseMimeType": "text/plain", # Request plain text for markdown output
            "temperature": 0.2, # Lower temperature for more focused output
            "topP": 0.9,
            "topK": 40
        }
    }

    headers = {'Content-Type': 'application/json'}

    print("Sending request to Gemini API...")
    try:
        response = requests.post(api_url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()  # Raise an HTTPError for bad responses (4xx or 5xx)
        result = response.json()

        if result.get("candidates") and result["candidates"][0].get("content") and \
           result["candidates"][0]["content"].get("parts") and \
           result["candidates"][0]["content"]["parts"][0].get("text"):
            markdown_table = result["candidates"][0]["content"]["parts"][0]["text"]
            return markdown_table
        else:
            return f"Gemini API did not return a valid response structure: {json.dumps(result, indent=2)}"
    except requests.exceptions.RequestException as e:
        return f"Request failed: {e}"
    except json.JSONDecodeError:
        return f"Failed to decode JSON response: {response.text}"
    except Exception as e:
        return f"An unexpected error occurred: {e}"

def analyze_ui_mockup_with_openai(image_path, HIGH_LEVEL_SCOPE):
    openai_api_key = os.getenv("OPENAI_API_KEY_SRS_GEN")
    if not openai_api_key:
        raise ValueError("OPENAI_API_KEY_SRS_GEN environment variable not set.")
    prompt_text = f"""
Analyze the provided UI screenshot thoroughly, considering the high-level project scope: {HIGH_LEVEL_SCOPE}. Your output should be structured into three distinct subsections:

Screen Name: First, identify the primary purpose or title of this UI screen and present it as a clear heading. Make this markdown bold line

Screen Description:
Provide a comprehensive description of the screen. This description should cover:
* The overall purpose of this module/screen to achieve the project's objective within the context of [INSERT_HIGH_LEVEL_SCOPE_HERE].
* What key functionalities or processes this specific module helps the user achieve.
* Any key workflows or user interactions facilitated by this screen.
Ensure this description is around 50-75 words and is plain text with no heading.

UI Elements Table:
Create a markdown table detailing only the key significant interactive elements, input fields, and links on the screen. For each identified element, include the following columns:
* Field Name: The explicit name, label, or clear identifier associated with the UI element.
* Type: The specific type of UI element (e.g., 'Text Input', 'Button', 'Dropdown', 'Text Area', 'Display Area', 'Link').
* Description: A concise explanation of its specific purpose and functionality within this UI context.
* Validation Rules: Any visible, implied, or standard validation rules or constraints applicable to input fields (e.g., 'Required', 'Numeric Only', 'File Types: .pdf, .txt', 'Specific Format').

Ensure the table is complete, accurately reflects only the key significant interactive elements - no need to mention self describing icons or texts , and strictly uses the specified column names.
Don't add any style or formatting to the table or headings.
The output should not include any unneccesary white or blank blocks.
    """
    base64_image = image_to_base64(image_path)
    if not base64_image:
        return "Failed to encode image."
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt_text},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}" if image_path.lower().endswith(('.jpg', '.jpeg')) else f"data:image/png;base64,{base64_image}"
                    }
                }
            ]
        }
    ]
    try:
        client = OpenAI(api_key=openai_api_key)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            temperature=0.95,
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return "Failed to analyze UI mockup with OpenAI."

def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

def process_mockup_zip_and_generate_markdown(zip_file_path, high_level_scope, use_openai=True):
    """
    Unzips the provided zip file, recursively analyzes all images, and returns combined markdown.
    """
    temp_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)
    
    print("temp_dir created by process_mockup_zip_and_generate_markdown " + temp_dir)

    def process_folder(folder_path, depth=1):
        entries = os.listdir(folder_path)
        folders = sorted([e for e in entries if os.path.isdir(os.path.join(folder_path, e))], key=natural_sort_key)
        files = sorted([e for e in entries if os.path.isfile(os.path.join(folder_path, e))], key=natural_sort_key)
        combined_markdown = ""
        # Skip heading for the top-level folder if it contains only folders
    # If this is the root or a top-level container (depth==1), skip heading and process subfolders directly
        if depth == 1:
            for subfolder in folders:
                subfolder_path = os.path.join(folder_path, subfolder)
                combined_markdown += process_folder(subfolder_path, depth + 1)
            # Also process files in the root if any (optional, or you can skip)
            for filename in files:
                if filename.startswith("._"):
                    continue
                if filename.lower().endswith((".png", ".jpg", ".jpeg")):
                    file_path = os.path.join(folder_path, filename)
                    print(f"Analyzing UI mockup from: {file_path}")
            return combined_markdown
        image_files = [filename for filename in files if filename.lower().endswith((".png", ".jpg", ".jpeg")) and not filename.startswith("._")]
        if image_files: 
            heading_level = '#' * (depth + 1)
            heading_text = os.path.basename(folder_path)
            heading_text = re.sub(r'^\d+_?', '', heading_text)
            heading_text = heading_text.replace('_', ' ').strip().title()
            combined_markdown += f"{heading_level} {heading_text}\n\n"
        
        # Process files in this folder
        for filename in files:
            if filename.startswith("._"):
                continue
            if filename.lower().endswith((".png", ".jpg", ".jpeg")):
                file_path = os.path.join(folder_path, filename)
                print(f"Analyzing UI mockup from: {file_path}")
                if use_openai:
                    markdown_output = analyze_ui_mockup_with_openai(file_path, high_level_scope)
                else:
                    markdown_output = analyze_ui_mockup_with_gemini(file_path, high_level_scope)
                screen_name_match = re.match(r"\*\*(.+?)\*\*\s*\n(.*?)(?:\n{2,}|$)", markdown_output, re.DOTALL)
                if screen_name_match:
                    screen_name = screen_name_match.group(1).strip()
                    screen_description = screen_name_match.group(2).strip()
                    combined_markdown += f"**{screen_name}**\n"
                    combined_markdown += f"![{os.path.basename(file_path)}]({file_path})\n\n"
                    combined_markdown += f"{screen_description}\n\n"
                    rest = markdown_output[screen_name_match.end():].strip()
                    if rest:
                        combined_markdown += rest + "\n\n"
                else:
                    combined_markdown += f"![{os.path.basename(file_path)}]({file_path})\n\n"
                    combined_markdown += markdown_output.strip() + "\n\n"
        for subfolder in folders:
            subfolder_path = os.path.join(folder_path, subfolder)
            combined_markdown += process_folder(subfolder_path, depth + 1)
        return combined_markdown

    combined_markdown = "## User Journey Flow\n\n"
    combined_markdown += process_folder(temp_dir, depth=1)
    return combined_markdown


def combined_srs_and_mockup_workflow(project_name, client_name, high_level_scope, transcript_list=None, artifact_list=None, image_zip_file=None, cloud_provider="Azure"):
    if not image_zip_file == None:
        # Save uploaded zip to a temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.zip') as tmp_zip:
            tmp_zip.write(image_zip_file.read())
            zip_path = tmp_zip.name
        
        # Analyze images and get markdown
        mockup_markdown = process_mockup_zip_and_generate_markdown(zip_path, high_level_scope, use_openai=True)

        print("===============mockup_markdown=======================")
        print(mockup_markdown)
    else:
        mockup_markdown=""

    # print(transcript_list)

    # Generate SRS markdown
    srs_markdown = generate_srs_text(project_name, client_name, high_level_scope, transcript_list, artifact_list, cloud_provider)
    print("===============srs_markdown=======================")
    print(srs_markdown)

    combined_markdown = srs_markdown + "\n\n" + mockup_markdown
    combined_markdown = combined_markdown.replace("```markdown", "")
    # Save combined markdown
    combined_md_path = f"outputs/{project_name}_SRS_and_Mockups.md"
    with open(combined_md_path, "w", encoding="utf-8") as f:
        f.write(combined_markdown)

    # Convert to docx
    output_docx_path = f"outputs/{project_name}_AI_Generated_SRS.docx"
    download_link = markdown_to_docx("s", combined_markdown, output_docx_path, project_name=project_name)
    return f"SRS Generation Complete 🎉 \nSRS and Mockup markdown saved to: {combined_md_path}\nCombined docx saved to: {output_docx_path}\nDownload link: {download_link}"
    

if __name__ == "__main__":

    # Set page config
    st.set_page_config(page_title="SRS Generator", layout="wide", page_icon="📝")

    # Custom CSS theme
    css_file_path = "custom_style.css"
    if os.path.exists(css_file_path):
        with open(css_file_path, "r") as css_file:
            custom_css = css_file.read()
        st.markdown(f"<style>{custom_css}</style>", unsafe_allow_html=True)
    else:
        st.warning(f"Custom CSS file '{css_file_path}' not found. Using default Streamlit styles.")


    # Page Header
    st.title("📝 Software Requirements Specification (SRS) Generator")
    st.markdown("Easily generate a professional-grade SRS document using transcripts, artifacts, and UI mockups.")

    # Form layout
    with st.form("srs_form"):
        col1, col2 = st.columns(2)

        with col1:
            project_name = st.text_input("🔧 Project Name*")
            client_name = st.text_input("👤 Client Name*")
            scope = st.text_area("📌 High-Level System Scope*", height=100)
            # Clean and define options properly
            platforms = ["AWS", "Azure", "GCP"]  
            cloud_platform = st.selectbox("☁️ Cloud Platform", options=platforms, index=None)
        with col2:
            transcripts = st.file_uploader("📄 Upload Transcript Files", type=["pdf", "txt", "docx"], accept_multiple_files=True)
            artifacts = st.file_uploader("📁 Upload Project Artifacts", type=["pdf", "txt", "docx", "xlsx", "png", "jpg", "jpeg"], accept_multiple_files=True)
            mockups = st.file_uploader("🖼️ Upload UI Mockup ZIP", type=["zip"])

        submitted = st.form_submit_button("🚀 Generate SRS")


    if submitted:
        # --- Validation Logic ---
        errors = []

        if not project_name:
            errors.append("Project Name is required.")
        if not client_name:
            errors.append("Client Name is required.")
        if not scope:
            errors.append("High-Level System Scope is required.")
        if cloud_platform is None: # Check if the 'None' placeholder is still selected
            cloud_platform = "Azure"

        # For example, if at least one transcript file is required:
        # if not upload_transcript_files:
        #     errors.append("At least one Transcript File is required.")
        transcripts_texts = []
        if transcripts:
            for uploaded_file in transcripts:
                file_name = uploaded_file.name
                file_type = uploaded_file.type # e.g., 'application/pdf', 'text/plain', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

                st.write(f"Processing file: {file_name} (Type: {file_type})")

                try:
                    if file_type == "application/pdf":
                        # For PDF: Use PyPDF2
                        pdf_reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
                        text = ""
                        for page_num in range(len(pdf_reader.pages)):
                            text += pdf_reader.pages[page_num].extract_text() or "" # extract_text() can return None
                        transcripts_texts.append(f"Content from {file_name}:\n{text}\n---END PDF---\n")

                    elif file_type == "text/plain":
                        # For TXT: Decode directly
                        text = uploaded_file.getvalue().decode("utf-8")
                        transcripts_texts.append(f"Content from {file_name}:\n{text}\n---END TXT---\n")

                    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        # For DOCX: Use python-docx
                        document = Document(io.BytesIO(uploaded_file.getvalue()))
                        text = ""
                        for paragraph in document.paragraphs:
                            text += paragraph.text + "\n"
                        transcripts_texts.append(f"Content from {file_name}:\n{text}\n---END DOCX---\n")

                    else:
                        st.warning(f"Unsupported file type for reading text: {file_name} ({file_type}). Skipping.")

                except Exception as e:
                    st.error(f"Error reading {file_name}: {e}")
                    errors.append(f"Error processing {file_name}.")


        artifacts_texts = []
        if transcripts:
            for uploaded_file in artifacts:
                file_name = uploaded_file.name
                file_type = uploaded_file.type # e.g., 'application/pdf', 'text/plain', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

                st.write(f"Processing file: {file_name} (Type: {file_type})")

                try:
                    if file_type == "application/pdf":
                        # For PDF: Use PyPDF2
                        pdf_reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
                        text = ""
                        for page_num in range(len(pdf_reader.pages)):
                            text += pdf_reader.pages[page_num].extract_text() or "" # extract_text() can return None
                        artifacts_texts.append(f"Content from {file_name}:\n{text}\n---END PDF---\n")
                    elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
                        # For Excel: Use pandas to read the file
                        excel_bytes = io.BytesIO(uploaded_file.getvalue())
                        try:
                            excel_file = pd.ExcelFile(excel_bytes)
                            for sheet_name in excel_file.sheet_names:
                                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                                text = df.to_csv(index=False, sep='\t')
                                artifacts_texts.append(f"Content from {file_name} (Sheet: {sheet_name}):\n{text}\n---END EXCEL SHEET---\n")
                        except Exception as e:
                            st.error(f"Error reading Excel file {file_name}: {e}")
                            errors.append(f"Error processing {file_name}.")
                    elif file_type == "text/plain":
                        # For TXT: Decode directly
                        text = uploaded_file.getvalue().decode("utf-8")
                        artifacts_texts.append(f"Content from {file_name}:\n{text}\n---END TXT---\n")

                    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        # For DOCX: Use python-docx
                        document = Document(io.BytesIO(uploaded_file.getvalue()))
                        text = ""
                        for paragraph in document.paragraphs:
                            text += paragraph.text + "\n"
                        artifacts_texts.append(f"Content from {file_name}:\n{text}\n---END DOCX---\n")

                    else:
                        st.warning(f"Unsupported file type for reading text: {file_name} ({file_type}). Skipping.")

                except Exception as e:
                    st.error(f"Error reading {file_name}: {e}")
                    errors.append(f"Error processing {file_name}.")


        if errors:
            for error in errors:
                st.error(error)
            st.warning("Please fill in all required fields.")
        else:
            # All fields are filled, proceed with processing
            st.success("Form submitted successfully! Processing SRS...")
            with st.spinner("Generating SRS document..."):
                result = combined_srs_and_mockup_workflow(
                    project_name, client_name, scope,
                    transcripts_texts, artifacts_texts, mockups, cloud_platform
                )
                st.success("✅ Done!")
                st.text_area("📬 Generation Status & File Paths", result, height=250)
